#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri May  6 15:04:11 2022

@author: doan
"""

from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
document = Document()

run = document.add_heading('Supplementary Material', level=0)

s = (
'This file includes figures, that provide the same information as those in the main text but for different runs (i.e., initializations).'
)


p = document.add_paragraph( s)

'''
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
'''


run = document.add_heading( level=1).add_run('Results for the WP test')
font = run.font
font.color.rgb = RGBColor(0x0, 0x0, 0)
#document.add_paragraph('Intense quote', style='Intense Quote')

'''
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
'''

#=== add figure 3


ia = 1
rr = range(1, 10)[:]
for ir in rr:
    print(ir)
    ifile = 'fig/fig03/SLP_DJF'+ '%.2d'%ir +'k04_rand.png'
    document.add_picture(ifile, width=Inches(6))
    

    s = ('Fig. S' +str(ia)+ \
         ' Similar with Figure 3 in the main manuscript but for the results from run ' +str(ir)+'. '
         'The winter SLP pattern revealed by S, C, E, and M k-means with k = 4. '
     '"H" indicates the location of the high and “L” indicates the location of the low. '
     'General silhouette analysis results are shown below the maps'
     ' where x-axis indicate the score, and y-axis the label of '
     'cluster numbered 1 – 4. Input data are ERA-Interim SLP data, '
     'which were re-gridded to Cartesian coordinates with a resolution '
     'of 200 x 200 km and grid size of 35 x 35. Daily data for December, '
     'January, and February over ten year 2005 – 2014 were used.'
    )
    
    run = document.add_paragraph( style='Caption').add_run(s)
    
    font = run.font
    font.color.rgb = RGBColor(0x0, 0x0, 0)
    ia = ia + 1
    
    

run = document.add_heading( level=1).add_run('Results for the CC test')
font = run.font
font.color.rgb = RGBColor(0x0, 0x0, 0)


for ir in rr:
    print(ir)
    #=== add figure 4
    print(ir)
    ifile = 'fig/fig04/AM_t_1950_y_k04r'+ '%.2d'%ir +'_rand.png'
    document.add_picture(ifile, width=Inches(6))

    s = ('Fig. S' +str(ia)+ 
         ' Similar with Figure 4 in the main manuscript but for the results from run ' +str(ir)+'. '
    'Result for CC experiment for cluster the climate change (temperature increase) time series over 134 weather stations over the entirety of Japan. Pattern were revealed by S, C, E, and M k-means with k = 4. Input data correspond to annual mean data collected over 70 years from 1951 – 2020 (subtracted by the mean of the first 30 years) and observed temperature achieved at in situ weather stations (dots in map) operated by the JMA. Time series of centroids and input vectors are shown in below panels together with general silhouette analysis results, where the x-axis indicates the score (S-score) and the y-axis presents the labels of clusters numbered 1 – 4. '
    )

    run = document.add_paragraph( style='Caption').add_run(s)
    
    font = run.font
    font.color.rgb = RGBColor(0x0, 0x0, 0)
    ia = ia + 1


run = document.add_heading( level=1).add_run('Results for the TC test')
font = run.font
font.color.rgb = RGBColor(0x0, 0x0, 0)


for ir in rr:
    print(ir)
    #=== add figure 5
    print(ir)
    ifile = 'fig/fig05/TC_ll_'+ '%.2d'%ir +'k04_rand.png'
    document.add_picture(ifile, width=Inches(6))
    

    s = ('Fig. S' +str(ia)+ ' Similar with Figure 5 in the maintext but for run ' +str(ir)+'. '
    'Results of the TC experiment for clustering tropical cyclone paths. The pattern was revealed by S, C, E, and M k-means, with k = 4. Input data are the best TC tracks obtained by the JMA from 1951 – 2020. Only TCs that passed the dashed box in the map are used to feed the k-means. Thus, a total of 863 TC tracking data points are used. The left side of each panel show the general silhouette analysis results, where the x-axis indicates the score (S-score) and y-axis presents the labels of clusters numbered 1 – 4. The centroid TC path is illustrated by the bold line, and the color is consistent with that in the silhouette diagram.'
    )
    
    
    run = document.add_paragraph( style='Caption').add_run(s)
    font = run.font
    font.color.rgb = RGBColor(0x0, 0x0, 0)
    
    ia = ia + 1










'''
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()
'''


odir = '../../../../../00\ Works/00\ Studies/00\ Papers-Lead-author/13\ kmeans/'
ofile = odir+'/SupplementaryInfo.docx'
document.save()
import os
os.system('cd '+odir+'; soffice --convert-to pdf SupplementaryInfo.docx')




